from docx import Document
import json
import os
import re
from datetime import datetime, timezone
from openai import OpenAI

client = OpenAI()

POLICIES_DIR = "policies"
OUTPUT_KB_PATH = "hr_kb.json"
OUTPUT_INDEX_PATH = "hr_kb_index.json"

# 檔名格式：HR-103-03_出勤管理辦法_202509.docx
#         QP-212-07_國內外出差管理辦法_202509.docx
FILENAME_RE = re.compile(r"^([A-Za-z]+-\d+-\d+)_([^_]+)_(\d{6})\.docx$")


def parse_filename(filename: str):
    base = os.path.basename(filename)
    m = FILENAME_RE.match(base)
    if not m:
        return {
            "policy_code": "未知版次",
            "policy_name": "未知辦法",
            "policy_month": "未知月份",
            "source_filename": base,
        }
    return {
        "policy_code": m.group(1),
        "policy_name": m.group(2),
        "policy_month": m.group(3),
        "source_filename": base,
    }


def read_docx_text(path: str) -> str:
    doc = Document(path)
    parts = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 表格（把每一列串成一行）
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line.strip():
                parts.append(line)

    return "\n".join(parts)


def chunk_text(text: str, chunk_size=650, overlap=120):
    """
    用字元長度切 chunk（對中文穩定），並保留 overlap
    """
    text = (text or "").strip()
    if not text:
        return []

    chunks = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + chunk_size, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks


def list_policy_docx_files():
    if not os.path.isdir(POLICIES_DIR):
        raise FileNotFoundError(f"找不到資料夾: ./{POLICIES_DIR}")

    files = []
    for name in os.listdir(POLICIES_DIR):
        if name.startswith("."):
            continue
        if not name.lower().endswith(".docx"):
            continue
        files.append(os.path.join(POLICIES_DIR, name))

    # 排序：月份新 -> 舊（其次檔名）
    def sort_key(p):
        meta = parse_filename(p)
        month = meta["policy_month"]
        # 未知月份放最後
        month_key = month if re.match(r"^\d{6}$", month) else "000000"
        return (month_key, meta["source_filename"])

    files.sort(key=sort_key, reverse=True)
    return files


def main():
    docx_files = list_policy_docx_files()
    if not docx_files:
        raise FileNotFoundError(f"./{POLICIES_DIR} 內找不到任何 .docx")

    generated_at = datetime.now(timezone.utc).isoformat()

    policies = []
    items = []

    for path in docx_files:
        meta = parse_filename(path)

        # 用 code+month 當 policy_id（方便 app.py 顯示/篩選）
        policy_id = f'{meta["policy_code"]}_{meta["policy_month"]}'
        meta["policy_id"] = policy_id
        meta["source_path"] = path.replace("\\", "/")
        meta["updated_at_utc"] = generated_at
        policies.append(meta)

        full_text = read_docx_text(path)
        chunks = chunk_text(full_text)

        print(f'📄 {meta["source_filename"]} -> {len(chunks)} chunks')

        for idx, ch in enumerate(chunks, start=1):
            emb = client.embeddings.create(
                model="text-embedding-3-small",
                input=ch
            )
            items.append({
                "policy_id": policy_id,
                "policy_code": meta["policy_code"],
                "policy_name": meta["policy_name"],
                "policy_month": meta["policy_month"],
                "source_filename": meta["source_filename"],
                "chunk_id": idx,
                "text": ch,
                "embedding": emb.data[0].embedding,
            })

    # 輸出 index
    index_out = {
        "generated_at_utc": generated_at,
        "policies_dir": POLICIES_DIR,
        "policies": policies,
    }
    with open(OUTPUT_INDEX_PATH, "w", encoding="utf-8") as f:
        json.dump(index_out, f, ensure_ascii=False, indent=2)

    # 輸出 kb
    kb_out = {
        "meta": {
            "generated_at_utc": generated_at,
            "policies_dir": POLICIES_DIR,
            "policy_count": len(policies),
            "item_count": len(items),
        },
        "policies": policies,   # 方便 app.py 直接取用
        "items": items,
    }
    with open(OUTPUT_KB_PATH, "w", encoding="utf-8") as f:
        json.dump(kb_out, f, ensure_ascii=False, indent=2)

    print(f"✅ 已輸出：{OUTPUT_INDEX_PATH}, {OUTPUT_KB_PATH}")


if __name__ == "__main__":
    main()
